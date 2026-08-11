from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "memoria_proyecto_intermodular.docx"
PDF_OUT = ROOT / "memoria_proyecto_intermodular.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "F2F4F7")
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Memoria del Proyecto Intermodular")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("One-page web de banco digital con PHP, MySQL, CSS modular y JavaScript").italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Proyecto: proyecto_intermodular\nFecha: junio de 2026\nAutor/a: [tu nombre]")

    doc.add_page_break()

    doc.add_heading("1. Introduccion", level=1)
    doc.add_paragraph(
        "Este documento presenta la memoria tecnica del proyecto intermodular. "
        "El proyecto consiste en una pagina web de tipo one-page inspirada en un banco digital, "
        "con una experiencia visual moderna, secciones reutilizables, contenido dinamico y una "
        "estructura preparada para evolucionar hacia herramientas financieras interactivas."
    )

    doc.add_heading("2. Objetivos", level=1)
    add_bullet(doc, "Diseñar una landing page completa para una entidad bancaria digital.")
    add_bullet(doc, "Organizar el proyecto mediante carpetas diferenciadas para codigo publico, configuracion, componentes y recursos.")
    add_bullet(doc, "Usar PHP para componer la pagina a partir de secciones reutilizables.")
    add_bullet(doc, "Cargar contenido de secciones desde base de datos MySQL y estructuras JSON.")
    add_bullet(doc, "Incorporar interactividad en cliente con JavaScript.")
    add_bullet(doc, "Mantener una base preparada para futuras calculadoras financieras y herramientas de seguimiento.")

    doc.add_heading("3. Tecnologias utilizadas", level=1)
    add_table(
        doc,
        ["Tecnologia", "Uso en el proyecto"],
        [
            ["HTML5", "Estructura semantica de la pagina y organizacion del contenido."],
            ["CSS3", "Estilos visuales, layouts, animaciones, variables, utilidades y media queries."],
            ["JavaScript", "Animaciones del hero y actualizacion dinamica de bloques mediante pestañas."],
            ["PHP", "Composicion de la pagina mediante includes y componentes reutilizables."],
            ["MySQL", "Almacenamiento de secciones dinamicas y datos asociados."],
            ["PDO", "Conexion a base de datos con manejo de errores y recuperacion de resultados."],
            ["JSON", "Formato de contenido para definir bloques estaticos y dinamicos de las secciones."],
            ["XAMPP", "Entorno local de ejecucion con Apache, PHP y MySQL."],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("4. Estructura del proyecto", level=1)
    add_table(
        doc,
        ["Carpeta o archivo", "Descripcion"],
        [
            ["public/", "Contiene la entrada principal de la aplicacion y los scripts JavaScript."],
            ["public/index.php", "Compone la pagina principal incluyendo las secciones PHP."],
            ["assets/css/", "Agrupa hojas de estilo separadas por responsabilidad: componentes, layouts, animaciones, media queries y utilidades."],
            ["assets/media/", "Contiene imagenes usadas en la landing page."],
            ["includes/", "Contiene las secciones que forman la pagina: hero, tarjetas, ahorro, seguridad, planes, enlaces y otras."],
            ["includes/components/", "Contiene funciones reutilizables para renderizar titulos, textos, botones, imagenes, planes, enlaces y controladores."],
            ["config/", "Contiene la conexion a base de datos, consultas y datos JSON de secciones."],
            ["config/sql/", "Scripts SQL para insertar o actualizar secciones dinamicas."],
            ["config/json/", "Archivos JSON que definen el contenido de las secciones."],
        ],
        [1.9, 4.6],
    )

    doc.add_heading("5. Arquitectura y funcionamiento", level=1)
    doc.add_paragraph(
        "La pagina principal se encuentra en public/index.php. Este archivo incluye el head, el header, "
        "las secciones principales dentro de main y las secciones del footer. Cada bloque de contenido "
        "se separa en un archivo PHP independiente dentro de includes, lo que facilita el mantenimiento "
        "y evita que toda la pagina dependa de un unico archivo grande."
    )
    doc.add_paragraph(
        "La carga de datos se inicia desde includes/head.php, que importa la conexion a base de datos, "
        "las consultas y los componentes reutilizables. En config/queries.php se recuperan las secciones "
        "activas desde la tabla secciones_dinamicas y se crea un mapa asociativo usando el html_id como clave. "
        "El contenido JSON de cada seccion se decodifica para poder usarlo desde PHP."
    )
    add_code(doc, "$web_data[$s['html_id']] = ['id' => $s['id'], 'nombre' => $s['nombre_interno'], 'contenido' => json_decode($s['contenido_json'], true)];")

    doc.add_heading("6. Componentes reutilizables", level=1)
    doc.add_paragraph(
        "El proyecto usa funciones de renderizado para evitar repetir codigo HTML. Por ejemplo, los componentes "
        "de titulo, texto, imagen, boton, enlace y plan reciben datos y clases CSS, y generan el HTML necesario. "
        "Esta idea hace que las secciones sean mas limpias y permite cambiar la forma de pintar un elemento desde "
        "un unico lugar."
    )
    add_bullet(doc, "renderTitle: genera titulos con nivel configurable.")
    add_bullet(doc, "renderText: muestra textos descriptivos a partir de datos estructurados.")
    add_bullet(doc, "renderImage: inserta imagenes usando los recursos de assets/media.")
    add_bullet(doc, "renderButton y renderLink: construyen llamadas a la accion y enlaces reutilizables.")
    add_bullet(doc, "renderPlan: permite mostrar planes de suscripcion de forma consistente.")

    doc.add_heading("7. Interactividad en JavaScript", level=1)
    doc.add_paragraph(
        "El archivo animations.js controla la animacion inicial del hero. Al detectar el desplazamiento de la rueda "
        "del raton, modifica transformaciones, opacidad y bloqueo temporal del scroll para crear una transicion "
        "entre estados visuales."
    )
    doc.add_paragraph(
        "El archivo sections.js añade interactividad a los botones de tipo pestaña. Cuando el usuario pulsa un boton, "
        "el script localiza la seccion correspondiente, lee los datos JSON embebidos y actualiza titulo, descripcion, "
        "imagen y boton sin recargar la pagina."
    )

    doc.add_heading("8. Diseño visual y estilos", level=1)
    doc.add_paragraph(
        "La parte visual se organiza en varios archivos CSS. Esta separacion permite mantener variables, estilos base, "
        "componentes, layouts, animaciones, utilidades y media queries en archivos independientes. El resultado es una "
        "base mas ordenada para un proyecto de tamaño medio."
    )
    add_bullet(doc, "variables.css define valores reutilizables del diseño.")
    add_bullet(doc, "components.css agrupa estilos de elementos reutilizables.")
    add_bullet(doc, "layout-hero.css y layout-general.css definen la disposicion de secciones.")
    add_bullet(doc, "animations.css contiene estilos vinculados a transiciones.")
    add_bullet(doc, "mediaqueries.css adapta la interfaz a diferentes tamaños de pantalla.")

    doc.add_heading("9. Base de datos", level=1)
    doc.add_paragraph(
        "La base de datos configurada se llama pi y la conexion se realiza desde config/db.php usando PDO. "
        "El proyecto recupera datos de una tabla de fotos y de una tabla de secciones dinamicas. La tabla "
        "secciones_dinamicas permite ordenar las secciones y almacenar el contenido en formato JSON."
    )
    add_table(
        doc,
        ["Elemento", "Funcion"],
        [
            ["config/db.php", "Configura host, base de datos, usuario, contraseña y opciones PDO."],
            ["config/queries.php", "Ejecuta consultas y prepara los datos para las secciones."],
            ["secciones_dinamicas", "Tabla usada para almacenar secciones y contenido JSON."],
            ["config/sql/", "Scripts de insercion y actualizacion de secciones dinamicas."],
        ],
        [2.0, 4.5],
    )

    doc.add_heading("10. Instalacion y ejecucion", level=1)
    add_number(doc, "Copiar la carpeta proyecto_intermodular dentro de htdocs de XAMPP.")
    add_number(doc, "Iniciar Apache y MySQL desde el panel de XAMPP.")
    add_number(doc, "Crear la base de datos pi en MySQL.")
    add_number(doc, "Ejecutar los scripts SQL incluidos en config/sql para cargar las secciones dinamicas.")
    add_number(doc, "Abrir en el navegador la ruta local correspondiente, por ejemplo http://localhost/proyectosDAM/proyecto_intermodular/public/.")

    doc.add_heading("11. Pruebas realizadas", level=1)
    add_bullet(doc, "Comprobacion de carga de la pagina principal desde public/index.php.")
    add_bullet(doc, "Revision de inclusion de secciones mediante archivos PHP independientes.")
    add_bullet(doc, "Comprobacion de conexion PDO y recuperacion de datos desde MySQL.")
    add_bullet(doc, "Revision de rutas de estilos e imagenes.")
    add_bullet(doc, "Comprobacion de animaciones e interacciones JavaScript en el navegador.")

    doc.add_heading("12. Mejoras futuras", level=1)
    add_bullet(doc, "Calculadora de gastos para que el usuario pueda clasificar y estimar sus movimientos.")
    add_bullet(doc, "Calculadora de ingresos con resumen mensual.")
    add_bullet(doc, "Calculadora de prestamos con cuotas, intereses y plazos.")
    add_bullet(doc, "Comparadora de cuentas bancarias.")
    add_bullet(doc, "Timeline de ingresos y gastos con previsiones y graficos.")
    add_bullet(doc, "Panel de administracion para editar secciones dinamicas desde una interfaz web.")

    doc.add_heading("13. Conclusion", level=1)
    doc.add_paragraph(
        "El proyecto cumple el objetivo de crear una landing page bancaria moderna y extensible. "
        "La solucion combina frontend, backend y base de datos, y aplica una estructura modular mediante includes, "
        "componentes reutilizables y datos JSON. Esta base permite mantener la web con mayor facilidad y ampliar "
        "el proyecto con nuevas herramientas financieras en futuras versiones."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Memoria del Proyecto Intermodular")

    doc.save(OUT)
    print(OUT)
    build_pdf()
    print(PDF_OUT)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="TitleMain",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=27,
        textColor=colors.HexColor("#0B2545"),
        spaceAfter=12,
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name="SubtitleMain",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=11,
        leading=15,
        alignment=1,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H1Custom",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=14,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="H2Custom",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=10,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="BodyCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="BulletCustom",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=14,
        leftIndent=16,
        firstLineIndent=-9,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        name="CodeCustom",
        parent=styles["BodyText"],
        fontName="Courier",
        fontSize=8.5,
        leading=11,
        leftIndent=18,
        spaceAfter=6,
    ))
    return styles


def p(text, styles):
    return Paragraph(text, styles["BodyCustom"])


def h1(text, styles):
    return Paragraph(text, styles["H1Custom"])


def bullet(text, styles):
    return Paragraph(f"&bull; {text}", styles["BulletCustom"])


def numbered(items, styles):
    return [Paragraph(f"{idx}. {text}", styles["BulletCustom"]) for idx, text in enumerate(items, 1)]


def pdf_table(headers, rows, widths, styles):
    data = [[Paragraph(f"<b>{h}</b>", styles["BodyCustom"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(cell), styles["BodyCustom"]) for cell in row])
    table = Table(data, colWidths=[w * inch for w in widths], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#000000")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C2CC")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def build_pdf():
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title="Memoria del Proyecto Intermodular",
    )
    story = [
        Spacer(1, 1.8 * inch),
        Paragraph("Memoria del Proyecto Intermodular", styles["TitleMain"]),
        Paragraph("One-page web de banco digital con PHP, MySQL, CSS modular y JavaScript", styles["SubtitleMain"]),
        Paragraph("Proyecto: proyecto_intermodular<br/>Fecha: junio de 2026<br/>Autor/a: [tu nombre]", styles["SubtitleMain"]),
        PageBreak(),
        h1("1. Introduccion", styles),
        p("Este documento presenta la memoria tecnica del proyecto intermodular. El proyecto consiste en una pagina web de tipo one-page inspirada en un banco digital, con una experiencia visual moderna, secciones reutilizables, contenido dinamico y una estructura preparada para evolucionar hacia herramientas financieras interactivas.", styles),
        h1("2. Objetivos", styles),
    ]
    for item in [
        "Diseñar una landing page completa para una entidad bancaria digital.",
        "Organizar el proyecto mediante carpetas diferenciadas para codigo publico, configuracion, componentes y recursos.",
        "Usar PHP para componer la pagina a partir de secciones reutilizables.",
        "Cargar contenido de secciones desde base de datos MySQL y estructuras JSON.",
        "Incorporar interactividad en cliente con JavaScript.",
        "Mantener una base preparada para futuras calculadoras financieras y herramientas de seguimiento.",
    ]:
        story.append(bullet(item, styles))

    story.extend([
        h1("3. Tecnologias utilizadas", styles),
        pdf_table(
            ["Tecnologia", "Uso en el proyecto"],
            [
                ["HTML5", "Estructura semantica de la pagina y organizacion del contenido."],
                ["CSS3", "Estilos visuales, layouts, animaciones, variables, utilidades y media queries."],
                ["JavaScript", "Animaciones del hero y actualizacion dinamica de bloques mediante pestañas."],
                ["PHP", "Composicion de la pagina mediante includes y componentes reutilizables."],
                ["MySQL", "Almacenamiento de secciones dinamicas y datos asociados."],
                ["PDO", "Conexion a base de datos con manejo de errores y recuperacion de resultados."],
                ["JSON", "Formato de contenido para definir bloques estaticos y dinamicos de las secciones."],
                ["XAMPP", "Entorno local de ejecucion con Apache, PHP y MySQL."],
            ],
            [1.55, 4.75],
            styles,
        ),
        Spacer(1, 8),
        h1("4. Estructura del proyecto", styles),
        pdf_table(
            ["Carpeta o archivo", "Descripcion"],
            [
                ["public/", "Contiene la entrada principal de la aplicacion y los scripts JavaScript."],
                ["public/index.php", "Compone la pagina principal incluyendo las secciones PHP."],
                ["assets/css/", "Agrupa hojas de estilo separadas por responsabilidad: componentes, layouts, animaciones, media queries y utilidades."],
                ["assets/media/", "Contiene imagenes usadas en la landing page."],
                ["includes/", "Contiene las secciones que forman la pagina: hero, tarjetas, ahorro, seguridad, planes, enlaces y otras."],
                ["includes/components/", "Contiene funciones reutilizables para renderizar titulos, textos, botones, imagenes, planes, enlaces y controladores."],
                ["config/", "Contiene la conexion a base de datos, consultas y datos JSON de secciones."],
                ["config/sql/", "Scripts SQL para insertar o actualizar secciones dinamicas."],
                ["config/json/", "Archivos JSON que definen el contenido de las secciones."],
            ],
            [1.7, 4.6],
            styles,
        ),
        Spacer(1, 8),
        h1("5. Arquitectura y funcionamiento", styles),
        p("La pagina principal se encuentra en public/index.php. Este archivo incluye el head, el header, las secciones principales dentro de main y las secciones del footer. Cada bloque de contenido se separa en un archivo PHP independiente dentro de includes, lo que facilita el mantenimiento y evita que toda la pagina dependa de un unico archivo grande.", styles),
        p("La carga de datos se inicia desde includes/head.php, que importa la conexion a base de datos, las consultas y los componentes reutilizables. En config/queries.php se recuperan las secciones activas desde la tabla secciones_dinamicas y se crea un mapa asociativo usando el html_id como clave.", styles),
        Paragraph("$web_data[$s['html_id']] = ['id' => $s['id'], 'nombre' => $s['nombre_interno'], 'contenido' => json_decode($s['contenido_json'], true)];", styles["CodeCustom"]),
        h1("6. Componentes reutilizables", styles),
        p("El proyecto usa funciones de renderizado para evitar repetir codigo HTML. Los componentes de titulo, texto, imagen, boton, enlace y plan reciben datos y clases CSS, y generan el HTML necesario.", styles),
    ])
    for item in [
        "renderTitle: genera titulos con nivel configurable.",
        "renderText: muestra textos descriptivos a partir de datos estructurados.",
        "renderImage: inserta imagenes usando los recursos de assets/media.",
        "renderButton y renderLink: construyen llamadas a la accion y enlaces reutilizables.",
        "renderPlan: permite mostrar planes de suscripcion de forma consistente.",
    ]:
        story.append(bullet(item, styles))

    story.extend([
        h1("7. Interactividad en JavaScript", styles),
        p("El archivo animations.js controla la animacion inicial del hero. Al detectar el desplazamiento de la rueda del raton, modifica transformaciones, opacidad y bloqueo temporal del scroll para crear una transicion entre estados visuales.", styles),
        p("El archivo sections.js añade interactividad a los botones de tipo pestaña. Cuando el usuario pulsa un boton, el script localiza la seccion correspondiente, lee los datos JSON embebidos y actualiza titulo, descripcion, imagen y boton sin recargar la pagina.", styles),
        h1("8. Diseño visual y estilos", styles),
        p("La parte visual se organiza en varios archivos CSS. Esta separacion permite mantener variables, estilos base, componentes, layouts, animaciones, utilidades y media queries en archivos independientes.", styles),
    ])
    for item in [
        "variables.css define valores reutilizables del diseño.",
        "components.css agrupa estilos de elementos reutilizables.",
        "layout-hero.css y layout-general.css definen la disposicion de secciones.",
        "animations.css contiene estilos vinculados a transiciones.",
        "mediaqueries.css adapta la interfaz a diferentes tamaños de pantalla.",
    ]:
        story.append(bullet(item, styles))

    story.extend([
        h1("9. Base de datos", styles),
        p("La base de datos configurada se llama pi y la conexion se realiza desde config/db.php usando PDO. El proyecto recupera datos de una tabla de fotos y de una tabla de secciones dinamicas. La tabla secciones_dinamicas permite ordenar las secciones y almacenar el contenido en formato JSON.", styles),
        pdf_table(
            ["Elemento", "Funcion"],
            [
                ["config/db.php", "Configura host, base de datos, usuario, contraseña y opciones PDO."],
                ["config/queries.php", "Ejecuta consultas y prepara los datos para las secciones."],
                ["secciones_dinamicas", "Tabla usada para almacenar secciones y contenido JSON."],
                ["config/sql/", "Scripts de insercion y actualizacion de secciones dinamicas."],
            ],
            [1.85, 4.45],
            styles,
        ),
        Spacer(1, 8),
        h1("10. Instalacion y ejecucion", styles),
    ])
    story.extend(numbered([
        "Copiar la carpeta proyecto_intermodular dentro de htdocs de XAMPP.",
        "Iniciar Apache y MySQL desde el panel de XAMPP.",
        "Crear la base de datos pi en MySQL.",
        "Ejecutar los scripts SQL incluidos en config/sql para cargar las secciones dinamicas.",
        "Abrir en el navegador la ruta local correspondiente, por ejemplo http://localhost/proyectosDAM/proyecto_intermodular/public/.",
    ], styles))

    story.append(h1("11. Pruebas realizadas", styles))
    for item in [
        "Comprobacion de carga de la pagina principal desde public/index.php.",
        "Revision de inclusion de secciones mediante archivos PHP independientes.",
        "Comprobacion de conexion PDO y recuperacion de datos desde MySQL.",
        "Revision de rutas de estilos e imagenes.",
        "Comprobacion de animaciones e interacciones JavaScript en el navegador.",
    ]:
        story.append(bullet(item, styles))

    story.append(h1("12. Mejoras futuras", styles))
    for item in [
        "Calculadora de gastos para que el usuario pueda clasificar y estimar sus movimientos.",
        "Calculadora de ingresos con resumen mensual.",
        "Calculadora de prestamos con cuotas, intereses y plazos.",
        "Comparadora de cuentas bancarias.",
        "Timeline de ingresos y gastos con previsiones y graficos.",
        "Panel de administracion para editar secciones dinamicas desde una interfaz web.",
    ]:
        story.append(bullet(item, styles))

    story.extend([
        h1("13. Conclusion", styles),
        p("El proyecto cumple el objetivo de crear una landing page bancaria moderna y extensible. La solucion combina frontend, backend y base de datos, y aplica una estructura modular mediante includes, componentes reutilizables y datos JSON. Esta base permite mantener la web con mayor facilidad y ampliar el proyecto con nuevas herramientas financieras en futuras versiones.", styles),
    ])

    doc.build(story)


if __name__ == "__main__":
    build()
